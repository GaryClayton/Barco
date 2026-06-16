from django.shortcuts import render

import os
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt

from companys.models import Company, CompanyContact, ClubContact, CRM
from events.models import Events, Contribution
from charity.models import Charity, CharityContact, CharityDonation
from pages.models import Page


def write_word_document(request):
    fn = os.path.join('barc_root/uploads/', 'full_database_output.docx')

    if os.path.exists(fn):
        try:
            os.remove(fn)
        except:
            return render(request, 'reports/doc_failure.html', {'Page_list': Page.objects.all()})

    # set-up base document
    document = Document()

    now = datetime.now()
    dt_string = now.strftime("%Y %m %d %H_%M")

    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = fn + " - saved at - " + dt_string

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    # Get all data - choice between hundreds of calls or 4
    comp = []
    comp_con = []
    comp_club = []
    comp_crm = []
    comp_cont = []
    comp_events = []
    char = []
    char_don = []
    char_con = []

    a = Company.objects.all()
    for c in a:
        temp = [c.id, c.name, c.sector, c.subsector, c.address, c.phone, c.email, c.web]
        comp.append(temp)
    b = CompanyContact.objects.all()
    for d in b:
        temp = [d.id, d.name, d.position, d.mobile, d.phone, d.email, str(d.company)]
        comp_con.append(temp)
    c = ClubContact.objects.all()
    for e in c:
        temp = [e.id, e.name, str(e.company)]
        comp_club.append(temp)
    dd = CRM.objects.all()
    for d in dd:
        temp = [d.id, str(d.username), str(d.created), str(d.company), d.comment]
        comp_crm.append(temp)

    ev = Events.objects.all()
    for d in ev:
        temp = [d.id, d.name, d.year, str(d.date)]
        comp_events.append(temp)

    cc = Contribution.objects.all()
    for c in cc:
        name = 'Unknown'
        for ev in comp_events:
            if str(c.event) == str(ev[0]):
                name = ev[1]
        temp = [c.id, c.support, c.value, str(name), str(c.company)]
        comp_cont.append(temp)

    comp.sort(key=lambda company: company[1])
    comp_crm.sort(key=lambda company: company[2], reverse=True)

    paragraph = document.add_heading('Bury St Edmunds Abbey Rotary Company Liaison Database', level=1)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph = document.add_heading('Company Data and All Associated Information', level=2)
    paragraph.paragraph_format.space_after = Pt(0)

    for row in comp:
        paragraph = document.add_paragraph('')
        paragraph.paragraph_format.space_after = Pt(0)

        db_line = str(row[1])
        paragraph = document.add_heading(db_line, level=3)
        paragraph.paragraph_format.space_after = Pt(0)
        db_line = "\t" + str(row[2]) + " \t" + str(row[3])
        paragraph = document.add_paragraph(db_line, style='Normal')
        paragraph.paragraph_format.space_after = Pt(0)
        db_line = "\t" + str(row[4]) + " \t" + str(row[5]) + " \t" + str(row[6]) + " \t" + str(row[7])
        paragraph = document.add_paragraph(db_line, style='Normal')
        paragraph.paragraph_format.space_after = Pt(0)

        paragraph = document.add_heading('\tCompany Points of Contact', level=4)
        paragraph.paragraph_format.space_after = Pt(0)
        for person in comp_con:  # Now add the contact points of authorised
            if str(person[6]) == str(row[0]):
                db_line = "\t\t" + str(person[1]) + " \t" + str(person[2]) + " \t" + str(person[3]) + " \t" \
                          + str(person[4]) + " \t" + str(person[5])
                paragraph = document.add_paragraph(db_line, style='Normal')
                paragraph.paragraph_format.space_after = Pt(0)

        paragraph = document.add_heading('\tBAR members responsible for liaison', level=4)
        paragraph.paragraph_format.space_after = Pt(0)
        for rep in comp_club:
            if str(rep[2]) == str(row[0]):  # is this a BAR representative for that company?
                db_line = "\t\t" + str(rep[1])
                paragraph = document.add_paragraph(db_line, style='Normal')
                paragraph.paragraph_format.space_after = Pt(0)

        paragraph = document.add_heading('\tSupport provided by Company', level=4)
        paragraph.paragraph_format.space_after = Pt(0)
        for rep in comp_cont:
            if str(rep[4]) == str(row[0]):  # is this a BAR representative for that company?
                db_line = "\t\t" + str(rep[3]) + " \t" + str(rep[1]) + " \t" + str(rep[2])
                paragraph = document.add_paragraph(db_line, style='Normal')
                paragraph.paragraph_format.space_after = Pt(0)

        paragraph = document.add_heading('\tCompany Relationship Contacts Loged', level=4)
        paragraph.paragraph_format.space_after = Pt(0)
        for rep in comp_crm:
            if str(rep[3]) == str(row[0]):
                db_line = "\t\t" + str(rep[1]) + " \t" + str(rep[2])
                paragraph = document.add_paragraph(db_line, style='Normal')
                paragraph.paragraph_format.space_after = Pt(0)
                db_line = "\t\t" + str(rep[4])
                paragraph = document.add_paragraph(db_line, style='Normal')
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph = document.add_paragraph('')
                paragraph.paragraph_format.space_after = Pt(0)

    # now compile and write all the charity data
    document.add_page_break()
    paragraph = document.add_heading('Charity Data and All Associated Information', level=2)
    paragraph.paragraph_format.space_after = Pt(0)

    x = Charity.objects.all()
    for d in x:
        temp = [d.id, d.name, d.sector, d.overview, d.web, d.address, d.phone]
        char.append(temp)

    y = CharityContact.objects.all()
    for d in y:
        temp = [d.id, d.name, d.title, d.phone, d.mobile, d.email, d.comment, str(d.charity)]
        char_con.append(temp)

    z = CharityDonation.objects.all()
    for d in z:
        temp = [d.id, str(d.date), d.value, d.comment, str(d.charity)]
        char_don.append(temp)

    # write charity data sorted by Charity
    paragraph = document.add_paragraph('')
    paragraph.paragraph_format.space_after = Pt(0)

    for row in char:
        db_line = str(row[1])
        paragraph = document.add_heading(db_line, level=3)  # Charity Name
        paragraph.paragraph_format.space_after = Pt(0)

        db_line = str(row[2])
        paragraph = document.add_paragraph(db_line, style='Normal')  # Other Charity information
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Pt(72)
        paragraph.paragraph_format.space_after = Pt(0)
        db_line = str(row[4]) + " \t" + str(row[5]) + " \t" + str(row[6])
        paragraph = document.add_paragraph(db_line, style='Normal')  # Other Charity information
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Pt(72)
        paragraph.paragraph_format.space_after = Pt(0)
        db_line = str(row[3])
        paragraph = document.add_paragraph(db_line, style='Normal')  # Charity Overview
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Pt(72)
        paragraph.paragraph_format.space_after = Pt(0)

        paragraph = document.add_heading('\tCharity Points of Contact', level=3)
        paragraph.paragraph_format.space_after = Pt(0)
        for person in comp_con:
            if str(person[6]) == str(row[0]):
                db_line = "\t\t" + str(person[1]) + " \t" + str(person[2]) + " \t" + str(person[3]) + " \t" \
                          + str(person[4]) + " \t" + str(person[5]) + " \t" + str(person[5])
                paragraph = document.add_paragraph(db_line, style='Normal')
                paragraph.paragraph_format.space_after = Pt(0)

        paragraph = document.add_heading('\tDonations given to Charity', level=3)
        paragraph.paragraph_format.space_after = Pt(0)
        for rep in char_don:
            if str(rep[4]) == str(row[0]):
                db_line = "\t\t" + str(rep[1]) + " \t" + str(rep[2]) + " \t" + str(rep[3])
                paragraph = document.add_paragraph(db_line, style='Normal')
                paragraph.paragraph_format.space_after = Pt(0)

    document.save(fn)

    return
